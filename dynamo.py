from transformers import pipeline
import pandas as pd
import torch
import pytesseract
from PIL import Image
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
import openpyxl
from bs4 import BeautifulSoup
import requests
import re

# Check if GPU is available
device = 0 if torch.cuda.is_available() else -1
print(f"Using device: {'GPU' if device == 0 else 'CPU'}")

# Load pre-trained GPT-2 model
generator = pipeline("text-generation", model="gpt2", device=device)

# Load finance-specific rules
finance_rules_df = pd.read_csv("finance_rules.csv")

# Function to adjust terminology based on dialect
def adjust_dialect(text, dialect="American"):
    if dialect == "British":
        replacements = {
            "color": "colour",
            "organize": "organise",
            "realize": "realise",
            "behavior": "behaviour",
            "center": "centre",
            "meter": "metre"
        }
        for us, uk in replacements.items():
            text = text.replace(us, uk)
    return text

# Function to extract text from various inputs
def extract_text(input_data, input_type="text"):
    if input_type == "text":
        return input_data
    elif input_type == "image":
        return pytesseract.image_to_string(Image.open(input_data))
    elif input_type == "pdf":
        with pdfplumber.open(input_data) as pdf:
            return " ".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif input_type == "docx":
        doc = Document(input_data)
        return " ".join(para.text for para in doc.paragraphs)
    elif input_type == "web":
        response = requests.get(input_data)
        soup = BeautifulSoup(response.text, "html.parser")
        return soup.get_text(separator=" ")
    elif input_type == "excel":
        wb = openpyxl.load_workbook(input_data)
        sheet = wb.active
        return " ".join(str(cell.value) for row in sheet.rows for cell in row if cell.value)
    return ""

# Function to classify requirements as functional or non-functional
def classify_requirements(text):
    functional = []
    non_functional = []
    sentences = re.split(r'[.!?]+', text)
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        if "must" in sentence.lower() or "shall" in sentence.lower():
            functional.append(sentence)
        elif any(keyword in sentence.lower() for keyword in ["performance", "security", "load", "under", "seconds", "compliance"]):
            non_functional.append(sentence)
        else:
            functional.append(sentence)
    
    return functional, non_functional

# Function to generate requirements
def generate_requirements(user_input, input_type="text", dialect="American"):
    # Extract text from input
    extracted_text = extract_text(user_input, input_type)
    
    # Generate initial requirements using GPT-2 with increased length
    prompt = f"Generate detailed software requirements for: {extracted_text}. Focus on finance applications. Provide at least 20 requirements."
    response = generator(prompt, max_length=500, num_return_sequences=1)[0]["generated_text"]
    
    # Add finance-specific rules from public domain knowledge
    relevant_rules = finance_rules_df["Rule"].tolist()
    requirements = response + "\n" + "\n".join(relevant_rules)
    
    # Adjust for dialect
    requirements = adjust_dialect(requirements, dialect)
    
    # Classify requirements
    functional, non_functional = classify_requirements(requirements)
    
    return functional, non_functional, requirements

# Function to calculate clarity score
def calculate_clarity_score(requirement_text):
    length = len(requirement_text.split())
    if length > 20:
        score = 8
        comment = "Good specificity"
    elif length > 10:
        score = 6
        comment = "Add more details"
    else:
        score = 4
        comment = "Too vague"
    return score, comment

# Function to generate Word document with standardized format (2-3 pages)
def generate_word_doc(functional, non_functional, filename="requirements.docx"):
    doc = Document()
    
    # Set document margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Title
    title = doc.add_heading("Software Requirements Specification", 0)
    title.alignment = 1  # Center
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.name = "Inter"
    
    # Introduction
    doc.add_heading("1. Introduction", level=1)
    intro = doc.add_paragraph(
        "This document outlines the requirements for the proposed system, generated by Dynamo: AI-Powered Requirement Writing. "
        "The system aims to address the needs of stakeholders by providing a comprehensive set of functional and non-functional requirements. "
        "The following sections detail the system's capabilities, ensuring alignment with industry standards and best practices."
    )
    for run in intro.runs:
        run.font.size = Pt(11)
        run.font.name = "Inter"
    
    # Purpose
    doc.add_heading("1.1 Purpose", level=2)
    purpose = doc.add_paragraph(
        "The purpose of this document is to define the requirements for a system that addresses the specified needs. "
        "It serves as a foundation for development, ensuring all stakeholders have a clear understanding of the system's functionality and constraints."
    )
    for run in purpose.runs:
        run.font.size = Pt(11)
        run.font.name = "Inter"
    
    # Scope
    doc.add_heading("1.2 Scope", level=2)
    scope = doc.add_paragraph(
        "This system will provide a robust solution for managing financial transactions, ensuring security, performance, and compliance with industry standards. "
        "It includes features for user authentication, transaction processing, and real-time monitoring, among others."
    )
    for run in scope.runs:
        run.font.size = Pt(11)
        run.font.name = "Inter"
    
    # Functional Requirements
    doc.add_heading("2. Functional Requirements", level=1)
    for i, req in enumerate(functional, 1):  # Include all functional requirements
        p = doc.add_paragraph(f"FR{i}: {req}", style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Inter"
    
    # Non-Functional Requirements
    doc.add_heading("3. Non-Functional Requirements", level=1)
    for i, req in enumerate(non_functional, 1):  # Include all non-functional requirements
        p = doc.add_paragraph(f"NFR{i}: {req}", style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Inter"
    
    # Assumptions and Constraints
    doc.add_heading("4. Assumptions and Constraints", level=1)
    assumptions = doc.add_paragraph(
        "4.1 Assumptions:\n"
        "- The system will have access to a stable internet connection for real-time monitoring.\n"
        "- Users are expected to have basic technical knowledge to interact with the system.\n\n"
        "4.2 Constraints:\n"
        "- The system must comply with PCI-DSS standards for payment processing.\n"
        "- Development must be completed within a 6-month timeline."
    )
    for run in assumptions.runs:
        run.font.size = Pt(11)
        run.font.name = "Inter"
    
    # Summary
    doc.add_heading("5. Summary", level=1)
    summary = doc.add_paragraph(
        "This SRS document provides a comprehensive overview of the system's requirements, ensuring clarity and alignment with stakeholder needs. "
        "It serves as a blueprint for development, testing, and deployment, facilitating a smooth project lifecycle."
    )
    for run in summary.runs:
        run.font.size = Pt(11)
        run.font.name = "Inter"
    
    doc.save(filename)
    return filename

# Function to extract user stories and export to Excel
def extract_user_stories(functional, filename="user_stories.xlsx"):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "User Stories"
    
    ws.append(["Summary", "Description", "Type", "Priority"])
    
    for i, req in enumerate(functional, 1):
        summary = f"User Story {i}"
        description = f"As a user, I want {req.lower()} so that I can achieve the desired functionality."
        ws.append([summary, description, "Story", "Medium"])
    
    wb.save(filename)
    return filename